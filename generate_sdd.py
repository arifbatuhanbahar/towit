# -*- coding: utf-8 -*-
"""
Towit Yazılım Tasarım Belgesi (SDD) Üretici
IEEE 1016-2009 uyumlu — Towit_SRS.docx ile aynı font/yapı
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

# ──────────────────────────────────────────────────────────────
# STİL SABİTLERİ  (SRS.docx'ten XML düzeyinde ölçüldü)
# ──────────────────────────────────────────────────────────────
FONT_HEADING = "Calibri"
FONT_BODY    = "Cambria"

COLOR_H1  = "366091"   # Heading 1 — koyu mavi
COLOR_H2  = "4F81BD"   # Heading 2/3 — orta mavi
COLOR_H3  = "4F81BD"

SZ_H1  = "28"          # 14 pt  (twips: sz × 2)
SZ_H2  = "26"          # 13 pt
SZ_H3  = "22"          # 11 pt
SZ_BODY = "22"         # 11 pt


# ──────────────────────────────────────────────────────────────
# XML YARDIMCILARI
# ──────────────────────────────────────────────────────────────

def _rpr_of_style(style):
    """Stil elementindeki w:rPr'yi bul ya da oluştur."""
    sd = style.element
    rPr = sd.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        sd.append(rPr)
    return rPr


def _set_rfonts(rPr, name):
    """w:rFonts'u 4 özniteliğiyle birlikte ayarla (SRS.docx birebir)."""
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def _set_sz(rPr, val):
    """w:sz ve w:szCs ayarla."""
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), val)


def _set_bold(rPr, val=True):
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), "1" if val else "0")


def _set_color(rPr, hex_val):
    el = rPr.find(qn("w:color"))
    if el is None:
        el = OxmlElement("w:color")
        rPr.append(el)
    el.set(qn("w:val"), hex_val)


def _set_spacing(style, before_twips=None, after_twips=None, line=None):
    """Stil düzeyinde paragraf aralığını ayarla."""
    sd  = style.element
    pPr = sd.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        sd.append(pPr)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    if before_twips is not None:
        sp.set(qn("w:before"), str(before_twips))
    if after_twips is not None:
        sp.set(qn("w:after"), str(after_twips))
    if line is not None:
        sp.set(qn("w:line"),     str(line))
        sp.set(qn("w:lineRule"), "auto")


# ──────────────────────────────────────────────────────────────
# STİL YAPILANDIRMASI
# ──────────────────────────────────────────────────────────────

def configure_styles(doc):
    """
    Tüm stilleri SRS.docx XML çıktısıyla birebir eşleşecek şekilde ayarla.
    Run düzeyinde override YOK — stil kalıbı cascade çalışır.
    """
    styles = doc.styles

    # ── docDefaults: Cambria 11pt, 1.15× satır, 10pt sonrası ──────────────
    styles_el   = styles.element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_defaults)

    rPrDef = doc_defaults.find(qn("w:rPrDefault"))
    if rPrDef is None:
        rPrDef = OxmlElement("w:rPrDefault")
        doc_defaults.append(rPrDef)
    rPr = rPrDef.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDef.append(rPr)
    _set_rfonts(rPr, FONT_BODY)
    _set_sz(rPr, SZ_BODY)

    pPrDef = doc_defaults.find(qn("w:pPrDefault"))
    if pPrDef is None:
        pPrDef = OxmlElement("w:pPrDefault")
        doc_defaults.append(pPrDef)
    pPr = pPrDef.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        pPrDef.append(pPr)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:after"),    "200")   # 10 pt
    sp.set(qn("w:line"),     "276")   # 1.15×
    sp.set(qn("w:lineRule"), "auto")

    # ── Heading 1: Calibri 14pt Bold #366091 ──────────────────────────────
    h1 = styles["Heading 1"]
    rPr1 = _rpr_of_style(h1)
    _set_rfonts(rPr1, FONT_HEADING)
    _set_bold(rPr1)
    _set_color(rPr1, COLOR_H1)
    _set_sz(rPr1, SZ_H1)
    _set_spacing(h1, before_twips=480, after_twips=100, line=276)

    # ── Heading 2: Calibri 13pt Bold #4F81BD ──────────────────────────────
    h2 = styles["Heading 2"]
    rPr2 = _rpr_of_style(h2)
    _set_rfonts(rPr2, FONT_HEADING)
    _set_bold(rPr2)
    _set_color(rPr2, COLOR_H2)
    _set_sz(rPr2, SZ_H2)
    _set_spacing(h2, before_twips=240, after_twips=80, line=276)

    # ── Heading 3: Calibri 11pt Bold #4F81BD ──────────────────────────────
    h3 = styles["Heading 3"]
    rPr3 = _rpr_of_style(h3)
    _set_rfonts(rPr3, FONT_HEADING)
    _set_bold(rPr3)
    _set_color(rPr3, COLOR_H3)
    _set_sz(rPr3, SZ_H3)
    _set_spacing(h3, before_twips=200, after_twips=60, line=276)

    # ── Normal stili ──────────────────────────────────────────────────────
    # Normal'in kendi rPr'si yoksa docDefaults devreye girer (Cambria 11pt)


# ──────────────────────────────────────────────────────────────
# PARAGRAF YARDIMCILARI  (run-level override YOK)
# ──────────────────────────────────────────────────────────────

def h1(doc, text):
    return doc.add_heading(text, level=1)

def h2(doc, text):
    return doc.add_heading(text, level=2)

def h3(doc, text):
    return doc.add_heading(text, level=3)

def body(doc, text):
    return doc.add_paragraph(text)

def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")

def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


# ──────────────────────────────────────────────────────────────
# TABLO  (SRS.docx: siyah kenarlık, bold başlık, renksiz)
# ──────────────────────────────────────────────────────────────

def _apply_table_borders(tbl):
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    bdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        bdr.append(b)
    tblPr.append(bdr)


def table(doc, headers, rows):
    """
    SRS.docx tablo görünümü:
    – Siyah tam kenarlık
    – Başlık satırı: bold, 10 pt
    – Veri satırları: 9 pt, ilk sütun bold
    """
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _apply_table_borders(tbl)

    # Başlık satırı
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    # Veri satırları
    for ri, row_vals in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_vals):
            cell = row.cells[ci]
            cell.text = ""
            p  = cell.paragraphs[0]
            r  = p.add_run(str(val))
            r.bold = (ci == 0)
            r.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)

    doc.add_paragraph()
    return tbl


# ──────────────────────────────────────────────────────────────
# DİYAGRAM BLOĞU
# ──────────────────────────────────────────────────────────────

def _set_para_shading(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def diagram(doc, title, mermaid_code):
    """
    Başlık kutusu (açık mavi) + Mermaid kaynak kodu (gri, Courier New).
    Word'de: ilgili kodu https://mermaid.live adresine yapıştırın.
    """
    doc.add_paragraph()

    # Başlık bandı
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_shading(p, "DEEAF1")
    r  = p.add_run(f"[ DİYAGRAM: {title} ]")
    r.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

    # Açıklama notu
    note = doc.add_paragraph()
    nr = note.add_run(
        "Not: Aşağıdaki Mermaid.js kodu, https://mermaid.live adresine "
        "yapıştırılarak görselleştirilebilir."
    )
    nr.italic = True
    nr.font.size = Pt(9)
    note.paragraph_format.space_after = Pt(0)

    # Kod satırları
    for line in mermaid_code.strip().split("\n"):
        cp = doc.add_paragraph()
        cr = cp.add_run(line if line.strip() else " ")
        cr.font.name = "Courier New"
        cr.font.size = Pt(9)
        _set_para_shading(cp, "F2F2F2")
        cp.paragraph_format.left_indent  = Pt(18)
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)

    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────
# KAPAK SAYFASI
# ──────────────────────────────────────────────────────────────

def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TOWIT")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x36, 0x60, 0x91)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Çekici Hizmetleri Platformu")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("YAZILIM TASARIM BELGESİ")
    r3.bold = True
    r3.font.size = Pt(22)
    r3.font.color.rgb = RGBColor(0x36, 0x60, 0x91)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Software Design Document  (SDD)")
    r4.font.size = Pt(14)
    r4.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for _ in range(3):
        doc.add_paragraph()

    table(doc, ["Bilgi", "Değer"], [
        ("Versiyon",    "1.0"),
        ("Tarih",       date.today().strftime("%d %B %Y")),
        ("Hazırlayan",  "Towit Yazılım Mimarlığı Ekibi"),
        ("Standart",    "IEEE Std 1016-2009 — Software Design Descriptions"),
        ("Kaynak SRS",  "Towit_SRS.docx  v1.0  (5 Nisan 2026)"),
        ("Gizlilik",    "Dahili — Dağıtımı Kısıtlıdır"),
    ])

    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# BÖLÜM 1 — GİRİŞ
# ──────────────────────────────────────────────────────────────

def section_1(doc):
    h1(doc, "1. Giriş")

    h2(doc, "1.1 Amaç")
    body(doc,
        "Bu Yazılım Tasarım Belgesi (SDD), Towit çekici hizmetleri platformunun teknik "
        "tasarım kararlarını ve mimari yapısını IEEE Std 1016-2009 standartlarına uygun "
        "biçimde belgelemektedir. Belge; yazılım mimarlarının, geliştiricilerin ve test "
        "mühendislerinin sistemi doğru ve tutarlı biçimde hayata geçirebilmesi için "
        "gereken tüm teknik referansı sağlamaktadır. Her tasarım bileşeni, "
        "Towit_SRS.docx belgesindeki ilgili gereksinimle (FR, UR, NFR, DR) "
        "eşleştirilmiştir."
    )

    h2(doc, "1.2 Kapsam")
    body(doc,
        "Towit, Türkiye pazarında faaliyet gösteren web tabanlı bir çekici hizmetleri "
        "aracılık platformudur. Sistem; müşteriler (araç sahipleri), operatörler "
        "(çekici şirketleri) ve harita servisi sağlayıcısı arasındaki iş akışını "
        "yöneten bir sunucu mimarisi ve tek sayfa uygulaması (SPA) biçiminde "
        "tasarlanmıştır."
    )
    body(doc, "Kapsam dahilindeki başlıca işlevler:")
    for item in [
        "Rol tabanlı kimlik doğrulama ve yetkilendirme (Müşteri / Operatör)",
        "Konum belirleme: GPS veya harita üzerinde manuel işaretleme",
        "Koordinata göre uygun operatör listeleme ve fiyat önizleme",
        "İş talebi oluşturma ve demo ödeme akışı",
        "Operatör tarafından talep kabul/red ve durum güncelleme",
        "OpenAPI ile versiyonlanmış REST API",
    ]:
        bullet(doc, item)

    h2(doc, "1.3 Tanımlar ve Kısaltmalar")
    table(doc, ["Terim / Kısaltma", "Tanım"], [
        ("Müşteri",          "Çekici hizmeti talep eden son kullanıcı"),
        ("Operatör",         "Kayıtlı çekici hizmeti sağlayıcısı (şahıs veya şirket)"),
        ("İş Talebi (Job)",  "Belirli teslim alma ve bırakma koordinatlarını içeren iş kaydı"),
        ("Tarife Kuralı",    "Baz ücret, km birim fiyatı, zaman çarpanlarından oluşan fiyat yapısı"),
        ("Önizleme Tutarı",  "Liste ekranında rota/mesafe varsayımlarıyla hesaplanan tahmini ücret"),
        ("Demo Ödeme",       "Gerçek PSP entegrasyonu olmaksızın ödeme adımını simüle eden uç nokta"),
        ("POI",              "İlgi Noktası — Yetkili servis, sanayi bölgesi vb."),
        ("REST",             "HTTP tabanlı web API mimari tarzı"),
        ("TLS",              "Şifrelenmiş ağ iletişim katmanı (HTTPS ile ilişkili)"),
        ("KVKK",             "6698 Sayılı Kişisel Verilerin Korunması Kanunu"),
        ("MVP",              "Minimum Uygulanabilir Ürün — ilk demo sürümü"),
        ("JWT",              "JSON Web Token — oturum yönetiminde kullanılan imzalı jeton"),
        ("SPA",              "Single Page Application — tek sayfa web uygulaması"),
        ("FSM",              "Sonlu Durum Makinesi — Finite State Machine"),
        ("RBAC",             "Rol Tabanlı Erişim Kontrolü — Role-Based Access Control"),
    ])

    h2(doc, "1.4 Referans Belgeler")
    for ref in [
        "[1] IEEE Std 1016-2009 — IEEE Standard for Information Technology — Systems Design — Software Design Descriptions",
        "[2] IEEE Std 830-1998 — IEEE Recommended Practice for Software Requirements Specifications",
        "[3] Towit_SRS.docx v1.0 — Towit Yazılım Gereksinimleri Spesifikasyonu (5 Nisan 2026)",
        "[4] 6698 Sayılı Kişisel Verilerin Korunması Kanunu ve ikincil düzenlemeler",
        "[5] Google Maps Platform Kullanım Şartları",
        "[6] OWASP ASVS — Application Security Verification Standard v4.0",
    ]:
        numbered(doc, ref)

    h2(doc, "1.5 Belge Yapısı")
    body(doc,
        "Belge şu bölümlerden oluşmaktadır: Bölüm 2 sistem mimarisine genel bakışı, "
        "Bölüm 3 IEEE 1016-2009 tasarım bakış açılarını, Bölüm 4 kullanım senaryosu "
        "tasarımını, Bölüm 5 veri tasarımını, Bölüm 6 bileşen tasarımını ve REST API "
        "kataloğunu, Bölüm 7 kullanıcı arayüzü tasarımını, Bölüm 8 güvenlik tasarımını "
        "ve Bölüm 9 gereksinim izlenebilirlik matrisini içermektedir."
    )
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# BÖLÜM 2 — SİSTEM MİMARİSİ
# ──────────────────────────────────────────────────────────────

def section_2(doc):
    h1(doc, "2. Sistem Mimarisine Genel Bakış")

    h2(doc, "2.1 Mimari Desen")
    body(doc,
        "Towit, klasik üç katmanlı (3-Tier) web mimarisi üzerine inşa edilmiştir. "
        "Sunum Katmanı (Presentation Layer) kullanıcı etkileşimini; İş Mantığı Katmanı "
        "(Business Logic Layer) tüm haritalama, fiyatlama ve iş akışı hesaplamalarını; "
        "Veri Katmanı (Data Layer) kalıcı depolamayı üstlenir. Tüm dış API çağrısı "
        "(harita, rota, POI) iş mantığı katmanında gerçekleştirilir; API anahtarları "
        "hiçbir zaman istemci tarafına açılmaz (NFR-03)."
    )

    h2(doc, "2.2 Teknoloji Yığını")
    table(doc, ["Katman", "Teknoloji / Araç", "Gerekçesi"], [
        ("Sunum — Frontend",    "React 18 + Vite + TypeScript",                  "Bileşen tabanlı SPA; hızlı yeniden oluşturma"),
        ("Harita (FE)",         "@vis.gl/react-google-maps",                     "Google Maps JS API React sarıcısı; NFR-06"),
        ("İş Mantığı — Backend","Node.js 20 LTS + TypeScript + Express",         "Takım yetkinliği; geniş ekosistem"),
        ("Kimlik Doğrulama",    "JWT RS256 + Refresh Token Rotasyonu",           "Durumsuz token; yatay ölçekleme"),
        ("Veritabanı",          "PostgreSQL 16",                                 "İlişkisel model; PostGIS koordinat desteği"),
        ("Oturum / Önbellek",   "Redis 7",                                       "Token iptali; Places API önbelleği (NFR-05)"),
        ("Harita API (BE)",     "Google Maps Platform — Directions, Places",     "Sunucu taraflı çağrı; anahtar gizliliği"),
        ("Konteyner",           "Docker + docker-compose",                       "Tekrarlanabilir geliştirme/üretim ortamı"),
        ("Altyapı",             "AWS ECS Fargate veya VPS",                      "Yatay ölçekleme (NFR-10)"),
        ("CI/CD",               "GitHub Actions",                                "Otomatik test ve dağıtım boru hattı"),
    ])

    h2(doc, "2.3 Bileşen Diyagramı")
    diagram(doc, "Bileşen Diyagramı — Towit Sistem Mimarisi", """graph TB
  subgraph Client["İstemci (Tarayıcı)"]
    FE["React SPA\n(Vite + TypeScript)"]
  end

  subgraph Server["Sunucu Katmanı (Node.js / Express)"]
    ROUTER["REST API Yönlendiricisi\n(/api/v1)"]
    AUTH["AuthService\n(JWT / Refresh Token)"]
    LOC["LocationService\n(Koordinat Doğrulama)"]
    PRICE["PricingService\n(Tarife Hesaplama)"]
    JOB["JobService\n(İş Yaşam Döngüsü)"]
    OP["OperatorService\n(Profil / Hizmet Alanı)"]
    POI["POIService\n(Places Proxy)"]
    CACHE["Redis\n(Önbellek / Token İptal)"]
  end

  subgraph Data["Veri Katmanı"]
    DB[("PostgreSQL 16")]
  end

  subgraph External["Dış Servisler"]
    MAPS["Google Maps Platform\n(Directions + Places)"]
  end

  FE -- "HTTPS / JSON" --> ROUTER
  ROUTER --> AUTH & LOC & PRICE & JOB & OP & POI
  AUTH  --> DB & CACHE
  LOC   --> DB
  PRICE --> DB & MAPS
  JOB   --> DB
  OP    --> DB
  POI   --> MAPS & CACHE""")

    h2(doc, "2.4 Dağıtım Mimarisi")
    body(doc,
        "Üretim ortamında frontend statik dosyaları bir CDN (CloudFront) üzerinden "
        "dağıtılır. Backend servisi durum bilgisi taşımayan konteynerler biçiminde "
        "AWS ECS Fargate üzerinde çalıştırılır; Application Load Balancer trafiği "
        "dengeler. PostgreSQL için AWS RDS Managed PostgreSQL, Redis için AWS "
        "ElastiCache kullanılır. Tüm servisler TLS 1.2+ ile korunan özel bir sanal "
        "ağda (VPC) birbirine bağlıdır (NFR-01). Yatay ölçekleme, yük dengeleyici "
        "arkasında birden fazla ECS görevi çalıştırılarak sağlanır (NFR-10)."
    )
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# BÖLÜM 3 — TASARIM BAKIŞ AÇILARI  (IEEE 1016-2009)
# ──────────────────────────────────────────────────────────────

def section_3(doc):
    h1(doc, "3. Tasarım Bakış Açıları (IEEE 1016-2009)")
    body(doc,
        "IEEE 1016-2009 standardı, bir yazılım tasarım belgesinin farklı paydaş "
        "perspektiflerinden incelenmesini önerir. Her Bakış Açısı (Viewpoint) için "
        "standart 'Tasarım Kaygıları (Design Concerns)' ve 'Tasarım Öğeleri "
        "(Design Elements)' alt bölümleri oluşturulmuştur."
    )

    # ── 3.1 Bağlam Bakış Açısı ────────────────────────────────
    h2(doc, "3.1 Bağlam Bakış Açısı (Context Viewpoint)")

    h3(doc, "3.1.1 Tasarım Kaygıları")
    body(doc,
        "Sistemin hangi dış aktörlerle, hangi protokoller üzerinden ve hangi "
        "veri sınırlarında etkileşime girdiğini tanımlamak."
    )

    h3(doc, "3.1.2 Tasarım Öğeleri")
    body(doc,
        "Sistem üç birincil dış aktörle etkileşir: Müşteri, Operatör ve Google Maps "
        "Platform. Müşteri ile Operatör platforma tarayıcı üzerinden HTTPS ile "
        "bağlanır. Google Maps Platform API'si yalnızca sunucu katmanından sorgulanır; "
        "istemci taraflı kodda API anahtarı yer almaz (NFR-03). Sistemin sınır "
        "koşulları: HTTPS dışında protokol kabul edilmez (NFR-01); gerçek PSP "
        "entegrasyonu MVP kapsamında dışarıda bırakılmıştır (FR-07); mobil native "
        "uygulama bu sürüm kapsamında yer almamaktadır."
    )

    # ── 3.2 Bileşim Bakış Açısı ───────────────────────────────
    h2(doc, "3.2 Bileşim Bakış Açısı (Composition Viewpoint)")

    h3(doc, "3.2.1 Tasarım Kaygıları")
    body(doc,
        "Sistemi bağımsız, tek sorumluluğu olan modüllere ayırmak ve her modülün "
        "sorumluluklarını netleştirmek."
    )

    h3(doc, "3.2.2 Tasarım Öğeleri")
    table(doc, ["Modül", "Sorumluluk", "SRS Referansı"], [
        ("AuthModule",            "Kayıt, giriş, token yenileme, oturum sonlandırma",                      "FR-01, FR-12"),
        ("LocationModule",        "Şehir seçimi, GPS/harita konum belirleme, koordinat doğrulama",         "FR-02, FR-13"),
        ("OperatorMatchModule",   "Koordinata göre aktif operatör filtreleme ve sıralama",                 "FR-03, FR-10"),
        ("PricingModule",         "Tarife kuralı getirme, rota/mesafe hesaplama, önizleme üretimi",        "FR-04, DR-01"),
        ("POIModule",             "Sunucu taraflı POI sorgulama ve Redis önbellekleme",                   "FR-05, NFR-05"),
        ("JobModule",             "İş talebi oluşturma, demo ödeme, FSM durum yönetimi",                  "FR-06–09, FR-14–15"),
        ("OperatorProfileModule", "Profil ve tarife yönetimi, hizmet alanı yapılandırması",               "FR-10, UR-07, UR-08"),
        ("SessionModule",         "JWT üretimi, refresh token rotasyonu, Redis token iptal kara listesi",  "FR-12, NFR-02"),
    ])

    # ── 3.3 Mantıksal Bakış Açısı ─────────────────────────────
    h2(doc, "3.3 Mantıksal Bakış Açısı (Logical Viewpoint)")

    h3(doc, "3.3.1 Tasarım Kaygıları")
    body(doc,
        "Sistem varlıkları arasındaki ilişkileri, sorumlulukları ve davranışsal "
        "sözleşmeleri tanımlamak."
    )

    h3(doc, "3.3.2 Tasarım Öğeleri")
    body(doc,
        "Detaylı sınıf diyagramı Bölüm 5.1'de sunulmaktadır. Temel ilişkiler: "
        "User ile OperatorProfile arasında 1:0..1 ilişkisi; OperatorProfile ile "
        "TariffRule arasında 1:N ilişkisi; her Job kaydı bir Müşteri (User) ve "
        "bir Operatör ile ilişkilendirilir; Session, User'a bağlı olup Redis'te "
        "TTL ile yönetilir."
    )

    # ── 3.4 Bağımlılık Bakış Açısı ────────────────────────────
    h2(doc, "3.4 Bağımlılık Bakış Açısı (Dependency Viewpoint)")

    h3(doc, "3.4.1 Tasarım Kaygıları")
    body(doc,
        "Dış paket ve servis bağımlılıklarını yönetmek; güvenlik açığı riskini "
        "en aza indirmek."
    )

    h3(doc, "3.4.2 Tasarım Öğeleri")
    table(doc, ["Paket", "Sürüm", "Amaç"], [
        ("express",                 "^4.18",   "HTTP sunucusu ve yönlendirme"),
        ("jsonwebtoken",            "^9.0",    "JWT üretimi ve doğrulama"),
        ("argon2",                  "^0.31",   "Parola hashleme — Argon2id (NFR-02)"),
        ("pg (node-postgres)",      "^8.11",   "PostgreSQL sürücüsü"),
        ("ioredis",                 "^5.3",    "Redis istemcisi"),
        ("zod",                     "^3.22",   "Girdi doğrulama ve şema tanımlama"),
        ("axios",                   "^1.6",    "Google Maps API HTTP istekleri (sunucu içi)"),
        ("react",                   "^18.2",   "Frontend UI kütüphanesi"),
        ("@vis.gl/react-google-maps","^1.x",   "Google Maps React entegrasyonu"),
        ("vite",                    "^5.x",    "Frontend derleme aracı"),
    ])

    # ── 3.5 Arayüz Bakış Açısı ────────────────────────────────
    h2(doc, "3.5 Arayüz Bakış Açısı (Interface Viewpoint)")

    h3(doc, "3.5.1 Tasarım Kaygıları")
    body(doc,
        "Sistemin dış dünyaya sunduğu arayüz sözleşmelerini tanımlamak; "
        "sürüm yönetimini güvence altına almak."
    )

    h3(doc, "3.5.2 Tasarım Öğeleri")
    body(doc,
        "Sistem dışarıya RESTful HTTP API arayüzü sunar. Tüm uç noktalar "
        "/api/v1 ön eki ile sürümlendirilmiştir (NFR-08). İstek ve yanıt "
        "gövdeleri JSON formatındadır. Kimlik doğrulama "
        "Authorization: Bearer <token> başlığı aracılığıyla sağlanır. "
        "Detaylı uç nokta kataloğu Bölüm 6.3'te sunulmaktadır. OpenAPI 3.1 "
        "spesifikasyonu geliştirme sürecinde ayrıca yayınlanacaktır."
    )

    # ── 3.6 Etkileşim Bakış Açısı ─────────────────────────────
    h2(doc, "3.6 Etkileşim Bakış Açısı (Interaction Viewpoint)")

    h3(doc, "3.6.1 Tasarım Kaygıları")
    body(doc,
        "Bileşenler arası mesaj akışını ve çağrı sırasını tanımlamak; "
        "veri tutarlılığını güvence altına almak."
    )

    h3(doc, "3.6.2 Tasarım Öğeleri")
    body(doc,
        "Ana iş akışı (müşteri talep oluşturma) için sekans diyagramı aşağıda "
        "sunulmaktadır. Bu akış FR-02, FR-03, FR-04, FR-06 ve FR-07 "
        "gereksinimlerinin birbirleriyle nasıl koordineli çalıştığını gösterir."
    )
    diagram(doc, "Sekans Diyagramı — Ana İş Akışı", """sequenceDiagram
  participant M  as Müşteri
  participant FE as Frontend (React SPA)
  participant BE as Backend (Node.js)
  participant DB as PostgreSQL
  participant GM as Google Maps API

  M->>FE: Şehir / Konum & Hedef Gir
  FE->>BE: GET /api/v1/operators/nearby?lat=&lng=&dest_lat=&dest_lng=
  BE->>GM: Directions API (sunucu taraflı — NFR-03)
  GM-->>BE: Rota ve Mesafe Bilgisi
  BE->>DB: Aktif operatörler + tarife sorgusu (FR-03, FR-10)
  DB-->>BE: Operatör listesi
  BE->>BE: Fiyat hesaplama — PricingService (FR-04)
  BE-->>FE: Operatörler + Önizleme Tutarları

  M->>FE: Operatör seç ve talebi onayla
  FE->>BE: POST /api/v1/jobs (FR-06)
  BE->>DB: Job oluştur — status: payment_pending
  BE-->>FE: Job ID döndü

  M->>FE: Demo ödemeyi onayla
  FE->>BE: POST /api/v1/jobs/:id/payment (FR-07)
  BE->>DB: status → open
  BE-->>FE: Ödeme başarılı

  Note over BE: Operatör bildirim mekanizması (polling / WS)
  M->>FE: Durum sorgula
  FE->>BE: GET /api/v1/jobs/:id (FR-09)
  BE-->>FE: Güncel durum (accepted / en_route)""")

    # ── 3.7 Durum Dinamiği Bakış Açısı ────────────────────────
    h2(doc, "3.7 Durum Dinamiği Bakış Açısı (State Dynamics Viewpoint)")

    h3(doc, "3.7.1 Tasarım Kaygıları")
    body(doc,
        "Job varlığının geçerli durum geçişlerini tanımlamak; "
        "izin verilmeyen geçişleri önlemek."
    )

    h3(doc, "3.7.2 Tasarım Öğeleri")
    body(doc,
        "Job FSM detayları Bölüm 5.3'te diyagram olarak sunulmaktadır. "
        "İzin verilmeyen geçişler backend tarafından HTTP 409 Conflict ile reddedilir."
    )
    table(doc, ["Kaynak Durum", "Hedef Durum", "Tetikleyen Olay / FR"], [
        ("payment_pending", "open",      "Demo ödeme başarılı — POST /jobs/:id/payment (FR-07)"),
        ("payment_pending", "cancelled", "Zaman aşımı veya müşteri iptali"),
        ("open",            "accepted",  "Operatör talebi kabul etti — PATCH /jobs/:id/status (FR-08)"),
        ("open",            "cancelled", "Operatör talebi reddetti — Alternatif operatör seçimi (FR-15)"),
        ("accepted",        "en_route",  "Operatör yola çıktı — PATCH /jobs/:id/status (FR-08)"),
        ("en_route",        "completed", "İş tamamlandı — PATCH /jobs/:id/status (FR-08)"),
        ("accepted",        "cancelled", "Sistem / operatör iptali"),
    ])
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# BÖLÜM 4 — KULLANIM SENARYOSU TASARIMI
# ──────────────────────────────────────────────────────────────

def section_4(doc):
    h1(doc, "4. Kullanım Senaryosu Tasarımı")
    body(doc,
        "Aşağıdaki Use Case diyagramı, sistemin Müşteri, Operatör ve Google Maps Platform "
        "aktörleriyle olan etkileşimlerini göstermektedir. Her kullanım senaryosu ilgili "
        "SRS gereksinimi ile ilişkilendirilmiştir."
    )
    diagram(doc, "Use Case Diyagramı — Towit Sistem Kullanım Senaryoları", """graph LR
  Musteri([Müşteri])
  Operator([Operatör])
  MapSvc([Google Maps API])

  subgraph Auth["Kimlik Yönetimi"]
    UC1["Kayıt Ol\n(FR-01)"]
    UC2["Giriş Yap\n(FR-01, FR-12)"]
  end

  subgraph Request["Talep Oluşturma Akışı"]
    UC3["Konum Belirle\n(FR-02, NFR-06)"]
    UC4["Hedef Seç\n(FR-05, UR-03)"]
    UC5["Operatör Listesi\n(FR-03, FR-04)"]
    UC6["Talep Oluştur\n(FR-06)"]
    UC7["Demo Ödeme\n(FR-07)"]
    UC8["Durum Takip\n(FR-09, UR-06)"]
  end

  subgraph OpPanel["Operatör Yönetimi"]
    UC9["Profil Düzenle\n(UR-07, FR-10)"]
    UC10["Tarife Yönet\n(UR-08)"]
    UC11["Kabul / Red\n(FR-08, FR-15)"]
    UC12["Durum Güncelle\n(FR-08, UR-10)"]
  end

  Musteri --> UC1 & UC2 & UC3 & UC4 & UC5 & UC6 & UC7 & UC8
  Operator --> UC2 & UC9 & UC10 & UC11 & UC12
  UC4 -->|Sunucu sorgusu| MapSvc
  UC5 -->|Rota / Mesafe API| MapSvc""")

    h2(doc, "4.1 Temel Kullanım Senaryoları")
    table(doc, ["ID", "Senaryo", "Açıklama", "FR Ref.", "Başarı Kriteri"], [
        ("UC-01", "Müşteri Kayıt ve Giriş",         "Müşteri sisteme kayıt olur ve giriş yapar",                          "FR-01, FR-12",    "JWT jeton döndürüldü; rol doğrulandı"),
        ("UC-02", "Konum ve Hedef Belirleme",        "GPS veya harita ile konum; metin/POI ile hedef belirleme",           "FR-02, FR-05",    "Koordinatlar doğrulandı; POI listesi gösterildi"),
        ("UC-03", "Operatör Listesi ve Fiyat",       "Sistem koordinata göre uygun operatörleri fiyatlarıyla listeler",   "FR-03, FR-04",    "En az bir operatör + önizleme tutarı gösterildi"),
        ("UC-04", "Talep Oluşturma ve Demo Ödeme",   "Müşteri operatör seçer, talep oluşturur, demo ödeme tamamlar",     "FR-06, FR-07",    "Job status: open; operatör bildirildi"),
        ("UC-05", "Operatör Talep Yanıtı",           "Operatör kabul veya reddeder; müşteri UI güncellenir",              "FR-08, FR-15",    "Job status güncel; alternatif operatör seçimi mümkün"),
        ("UC-06", "İş Durumu Güncelleme",            "Operatör durumu en_route → completed olarak günceller",            "FR-08, UR-10",    "Job status: completed; kayıt kapandı"),
    ])
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# BÖLÜM 5 — VERİ TASARIMI
# ──────────────────────────────────────────────────────────────

def section_5(doc):
    h1(doc, "5. Veri Tasarımı")

    h2(doc, "5.1 Sınıf ve Varlık-İlişki Diyagramı")
    diagram(doc, "Sınıf Diyagramı — Veri Modeli", """classDiagram
  class User {
    +UUID       id          PK
    +String     email       UNIQUE NOT NULL
    +String     passwordHash NOT NULL
    +Enum       role        customer | operator
    +Boolean    isActive
    +DateTime   createdAt
    +DateTime   updatedAt
    +login()    Boolean
    +refresh()  String
  }

  class OperatorProfile {
    +UUID    id              PK
    +UUID    userId          FK UNIQUE
    +String  businessName    NOT NULL
    +String  vehiclePlate
    +Float   serviceAreaLat
    +Float   serviceAreaLng
    +Float   serviceAreaRadiusKm
    +Boolean isActive
  }

  class TariffRule {
    +UUID    id                 PK
    +UUID    operatorId         FK
    +Float   baseFare
    +Float   perKmRate
    +Float   nightMultiplier
    +Float   holidayMultiplier
    +Date    effectiveFrom
    +Date    effectiveTo
    +calculatePreview(distKm, hour) Float
  }

  class Job {
    +UUID    id            PK
    +UUID    customerId    FK
    +UUID    operatorId    FK
    +Float   pickupLat
    +Float   pickupLng
    +Float   dropoffLat
    +Float   dropoffLng
    +String  pickupAddress
    +String  dropoffAddress
    +Float   priceSnapshot
    +Float   distanceKm
    +Enum    status
    +DateTime createdAt
    +updateStatus(s) void
  }

  class Session {
    +UUID     id
    +UUID     userId        FK
    +String   refreshToken  UNIQUE
    +DateTime expiresAt
    +Boolean  isRevoked
    +revoke() void
  }

  User        "1" --> "0..1"  OperatorProfile : sahiptir
  OperatorProfile "1" --> "1..*" TariffRule  : tanımlar
  User        "1" --> "0..*"  Job            : oluşturur (müşteri)
  OperatorProfile "1" --> "0..*" Job          : atanır (operatör)
  User        "1" --> "0..*"  Session        : sahiptir""")

    h2(doc, "5.2 Veritabanı Şeması")
    body(doc, "Aşağıdaki tablolar PostgreSQL 16 veritabanında tanımlanacaktır.")

    schemas = [
        ("users", [
            ("id",           "UUID",         "PRIMARY KEY  DEFAULT gen_random_uuid()"),
            ("email",        "VARCHAR(255)", "UNIQUE NOT NULL"),
            ("password_hash","TEXT",         "NOT NULL"),
            ("role",         "VARCHAR(20)",  "NOT NULL  CHECK (role IN ('customer','operator'))"),
            ("is_active",    "BOOLEAN",      "DEFAULT TRUE"),
            ("created_at",   "TIMESTAMPTZ",  "DEFAULT NOW()"),
            ("updated_at",   "TIMESTAMPTZ",  "DEFAULT NOW()"),
        ]),
        ("operator_profiles", [
            ("id",                    "UUID",        "PRIMARY KEY  DEFAULT gen_random_uuid()"),
            ("user_id",               "UUID",        "UNIQUE  REFERENCES users(id)  ON DELETE CASCADE"),
            ("business_name",         "VARCHAR(255)","NOT NULL"),
            ("vehicle_plate",         "VARCHAR(20)", ""),
            ("service_area_lat",      "FLOAT",       ""),
            ("service_area_lng",      "FLOAT",       ""),
            ("service_area_radius_km","FLOAT",       "DEFAULT 50"),
            ("is_active",             "BOOLEAN",     "DEFAULT TRUE"),
        ]),
        ("tariff_rules", [
            ("id",                  "UUID",          "PRIMARY KEY  DEFAULT gen_random_uuid()"),
            ("operator_id",         "UUID",          "REFERENCES operator_profiles(id)  ON DELETE CASCADE"),
            ("base_fare",           "NUMERIC(10,2)", "NOT NULL"),
            ("per_km_rate",         "NUMERIC(10,2)", "NOT NULL"),
            ("night_multiplier",    "NUMERIC(4,2)",  "DEFAULT 1.5"),
            ("holiday_multiplier",  "NUMERIC(4,2)",  "DEFAULT 2.0"),
            ("effective_from",      "DATE",          "NOT NULL"),
            ("effective_to",        "DATE",          "NULLABLE"),
        ]),
        ("jobs", [
            ("id",             "UUID",          "PRIMARY KEY  DEFAULT gen_random_uuid()"),
            ("customer_id",    "UUID",          "REFERENCES users(id)"),
            ("operator_id",    "UUID",          "REFERENCES operator_profiles(id)"),
            ("pickup_lat",     "FLOAT",         "NOT NULL"),
            ("pickup_lng",     "FLOAT",         "NOT NULL"),
            ("dropoff_lat",    "FLOAT",         "NOT NULL"),
            ("dropoff_lng",    "FLOAT",         "NOT NULL"),
            ("pickup_address", "TEXT",          ""),
            ("dropoff_address","TEXT",          ""),
            ("price_snapshot", "NUMERIC(10,2)", "NOT NULL"),
            ("distance_km",    "FLOAT",         ""),
            ("status",         "VARCHAR(30)",   "DEFAULT 'payment_pending'"),
            ("created_at",     "TIMESTAMPTZ",   "DEFAULT NOW()"),
        ]),
        ("sessions", [
            ("id",            "UUID",       "PRIMARY KEY  DEFAULT gen_random_uuid()"),
            ("user_id",       "UUID",       "REFERENCES users(id)  ON DELETE CASCADE"),
            ("refresh_token", "TEXT",       "UNIQUE NOT NULL"),
            ("expires_at",    "TIMESTAMPTZ","NOT NULL"),
            ("ip_address",    "INET",       ""),
            ("is_revoked",    "BOOLEAN",    "DEFAULT FALSE"),
        ]),
    ]
    for tname, cols in schemas:
        p = doc.add_paragraph()
        r = p.add_run(f"Tablo: {tname}")
        r.bold = True
        table(doc, ["Sütun Adı", "Veri Tipi", "Kısıtlama / Açıklama"], cols)

    h2(doc, "5.3 İş Yaşam Döngüsü — Durum Makinesi")
    diagram(doc, "Durum Makinesi Diyagramı — Job FSM", """stateDiagram-v2
  [*] --> payment_pending : Job oluşturuldu (FR-06)

  payment_pending --> open      : Demo ödeme başarılı\nPOST /jobs/:id/payment (FR-07)
  payment_pending --> cancelled : Zaman aşımı (TTL 15 dk.) veya müşteri iptali

  open --> accepted  : Operatör kabul etti\nPATCH /jobs/:id/status (FR-08)
  open --> cancelled : Operatör reddetti — Alternatif operatör seçimi (FR-15)

  accepted --> en_route  : Operatör yola çıktı\nPATCH /jobs/:id/status (FR-08, UR-10)
  accepted --> cancelled : Sistem / operatör iptali

  en_route --> completed : İş tamamlandı\nPATCH /jobs/:id/status (FR-08, UR-10)

  completed --> [*]
  cancelled --> [*]

  note right of payment_pending : TTL süresi aşılırsa sistem otomatik iptal eder.
  note right of open            : Eşzamanlı açık talep politikası ürün kararıyla belirlenir (FR-14).""")
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# BÖLÜM 6 — BİLEŞEN TASARIMI VE API KATALOĞU
# ──────────────────────────────────────────────────────────────

def section_6(doc):
    h1(doc, "6. Bileşen Tasarımı")

    h2(doc, "6.1 Frontend Modülleri")
    table(doc, ["Modül", "Sorumluluk", "FR / UR"], [
        ("AuthModule",              "Giriş, kayıt, token yenileme, korunan rota HOC",                       "FR-01, FR-12"),
        ("LocationModule",          "GPS izin isteği, harita konum seçimi, şehir dropdown",                 "FR-02, NFR-06"),
        ("POIModule",               "Hedef adres metin arama, POI kategorisi listesi",                      "FR-05, UR-03"),
        ("OperatorListModule",      "Operatör kartı listesi, fiyat önizleme, sıralama",                     "FR-03, FR-04"),
        ("JobCreateModule",         "Operatör seçimi, talep form onayı, demo ödeme akışı",                 "FR-06, FR-07"),
        ("JobStatusModule",         "Aktif talep durumu gösterimi; polling veya WebSocket",                 "FR-09, FR-15"),
        ("OperatorDashboardModule", "Profil düzenleme, tarife yönetimi, gelen talepler, durum güncelleme", "UR-07, UR-08, FR-08"),
    ])

    h2(doc, "6.2 Backend Servisleri")
    table(doc, ["Servis", "Sorumluluk", "FR / NFR"], [
        ("AuthService",         "JWT üretimi (RS256), refresh token rotasyonu, Redis iptal kara listesi",         "FR-01, FR-12, NFR-02"),
        ("LocationService",     "Şehir / koordinat doğrulama, sınır kontrolü, anlamlı hata kodu üretimi",        "FR-02, FR-13"),
        ("OperatorMatchService","Koordinata göre aktif operatör filtreleme (mesafe veya PostGIS polygon)",         "FR-03, FR-10"),
        ("PricingService",      "Tarife kuralı getirme, Directions API mesafe çağrısı, önizleme hesaplama",      "FR-04, DR-01, NFR-04"),
        ("POIService",          "Google Places API sorgulama, Redis önbellekleme (TTL 3600 s)",                  "FR-05, NFR-05"),
        ("JobService",          "İş oluşturma, FSM geçiş doğrulama, durum güncelleme, eşzamanlı talep politikası","FR-06–09, FR-14, FR-15"),
        ("OperatorService",     "Profil CRUD, tarife CRUD, hizmet alanı yapılandırması",                         "FR-10, UR-07, UR-08"),
        ("DemoPaymentService",  "Sahte ödeme uç noktası — gerçek PSP entegrasyonu olmaksızın success yanıtı",    "FR-07"),
    ])

    h2(doc, "6.3 REST API Uç Nokta Kataloğu")
    body(doc,
        "Tüm uç noktalar /api/v1 ön eki taşır. Yetkilendirme gerektiren uç noktalar "
        "Authorization: Bearer <access_token> başlığını zorunlu kılar. OpenAPI 3.1 "
        "spesifikasyonu geliştirme sürecinde ayrıca yayınlanacaktır (NFR-08)."
    )
    table(doc,
        ["Metot", "Yol  (/api/v1)", "Yetki", "Açıklama", "FR Ref."],
        [
            ("POST",  "/auth/register",          "Açık",      "Müşteri / operatör kayıt oluşturma",                  "FR-01"),
            ("POST",  "/auth/login",             "Açık",      "Giriş; access + refresh token döndürme",              "FR-01"),
            ("POST",  "/auth/refresh",           "Açık",      "Refresh token ile yeni access token",                 "FR-12"),
            ("POST",  "/auth/logout",            "Kimlik D.", "Token iptal ve oturum sonlandırma",                   "FR-12"),
            ("GET",   "/operators/nearby",       "Kimlik D.", "Koordinata göre aktif operatörler + önizleme tutarı", "FR-03, FR-04"),
            ("GET",   "/operators/:id",          "Kimlik D.", "Tek operatör profil detayı",                          "FR-10"),
            ("PUT",   "/operators/:id/profile",  "Operatör",  "Operatör profil güncelleme",                          "UR-07, FR-10"),
            ("GET",   "/operators/:id/tariff",   "Kimlik D.", "Operatör tarife kuralı getirme",                      "FR-04"),
            ("PUT",   "/operators/:id/tariff",   "Operatör",  "Tarife kuralı güncelleme",                            "UR-08"),
            ("GET",   "/places/search",          "Kimlik D.", "Adres / POI metin arama (Places API proxy)",          "FR-05"),
            ("POST",  "/jobs",                   "Müşteri",   "Yeni iş talebi oluşturma (status: payment_pending)",  "FR-06"),
            ("GET",   "/jobs",                   "Kimlik D.", "Rol bazlı iş listesi",                                "FR-09"),
            ("GET",   "/jobs/:id",               "Kimlik D.", "Tek iş detayı ve güncel durum",                       "FR-09"),
            ("POST",  "/jobs/:id/payment",       "Müşteri",   "Demo ödeme onaylama (status: open)",                  "FR-07"),
            ("PATCH", "/jobs/:id/status",        "Operatör",  "İş durumu güncelleme (FSM geçişi)",                   "FR-08, FR-15"),
        ]
    )
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# BÖLÜM 7 — KULLANICI ARAYÜZÜ TASARIMI
# ──────────────────────────────────────────────────────────────

def section_7(doc):
    h1(doc, "7. Kullanıcı Arayüzü Tasarımı")
    body(doc,
        "Kullanıcı arayüzü tek bir React SPA olarak tasarlanmıştır. Mobil ve masaüstü "
        "tarayıcıları destekleyen duyarlı (responsive) tasarım uygulanmaktadır. Dil "
        "Türkçe olup tüm hata mesajları teknik jargon içermeksizin anlaşılır biçimde "
        "yazılmıştır (UR-01, NFR-06)."
    )

    h2(doc, "7.1 Müşteri Navigasyon Akışı")
    table(doc, ["Adım", "Ekran", "İçerik / Eylem"], [
        ("1", "Karşılama / Giriş",  "E-posta + şifre girişi; kayıt bağlantısı; JWT alınır"),
        ("2", "Konum Belirleme",    "GPS butonu veya harita işaretleme; şehir dropdown zorunlu (FR-02)"),
        ("3", "Hedef Seçimi",       "Metin arama veya POI kategorisi; Places API sonuçları (FR-05)"),
        ("4", "Operatör Listesi",   "Fiyat + mesafeye göre sıralama; kart: isim, araç, önizleme tutar"),
        ("5", "Talep Onay",         "Seçilen operatör, adresler, tahmini ücret özeti"),
        ("6", "Demo Ödeme",         "Tutar gösterimi; 'Ödemeyi Tamamla' düğmesi; gerçek kart alınmaz"),
        ("7", "Durum Takip",        "Gerçek zamanlı: bekleniyor / kabul edildi / yolda / tamamlandı"),
    ])

    h2(doc, "7.2 Operatör Navigasyon Akışı")
    table(doc, ["Adım", "Ekran", "İçerik / Eylem"], [
        ("1", "Giriş",               "Operatör e-posta + şifre ile giriş yapar"),
        ("2", "Operatör Paneli",     "Bekleyen talepler listesi; profil aktif/pasif durumu"),
        ("3", "Talep Detay",         "Konum, hedef, önizleme tutarı; Kabul / Red düğmeleri (FR-08)"),
        ("4", "Aktif İş",            "Yola Çıktım / İş Tamamlandı düğmeleri (FR-08, UR-10)"),
        ("5", "Profil Düzenleme",    "İşletme bilgileri, araç plakası, hizmet alanı (FR-10, UR-07)"),
        ("6", "Tarife Yönetimi",     "Baz ücret, km birim fiyatı, gece / tatil çarpanları (UR-08)"),
    ])

    h2(doc, "7.3 Erişilebilirlik ve Kullanılabilirlik")
    body(doc,
        "Form alanları ARIA etiketleri ile işaretlenir; renk tek başına bilgi taşıyıcı "
        "olarak kullanılamaz (NFR-07). Birincil senaryo (konum, hedef, liste, talep) "
        "tek el ile tamamlanabilir biçimde tasarlanır (NFR-06). Dokunmatik alanların "
        "minimum boyutu 44×44 px olarak belirlenir (WCAG 2.1 AA). Hata mesajları form "
        "alanı yanında, anında ve Türkçe olarak gösterilir."
    )
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# BÖLÜM 8 — GÜVENLİK TASARIMI
# ──────────────────────────────────────────────────────────────

def section_8(doc):
    h1(doc, "8. Güvenlik Tasarımı")
    body(doc,
        "Aşağıdaki kontroller, SRS.docx'teki NFR ve DR gereksinimleri ile "
        "OWASP ASVS v4.0 kılavuzu doğrultusunda tasarlanmıştır."
    )
    table(doc, ["Kontrol / Referans", "Tasarım Kararı"], [
        ("NFR-01 — TLS Şifrelemesi",
         "Üretim ortamında tüm HTTP trafiği TLS 1.2+ ile şifrelenir. Zayıf "
         "şifreler (RC4, 3DES) ve eski TLS sürümleri devre dışı bırakılır. "
         "HSTS başlığı aktifleştirilir (max-age ≥ 31 536 000)."),
        ("NFR-02 — Parola Güvenliği",
         "Kullanıcı parolaları Argon2id algoritması ile hashlenir (m=65536, t=3, p=4). "
         "Düz metin parola hiçbir koşulda veritabanına veya log sistemine yazılmaz."),
        ("NFR-03 — API Anahtar Yönetimi",
         "Google Maps Platform API anahtarları yalnızca sunucu ortam değişkenlerinde "
         "(environment variables / AWS Secrets Manager) saklanır. İstemci taraflı "
         "kodda veya git geçmişinde yer almaz."),
        ("NFR-05 — Önbellekleme",
         "Places API sonuçları Redis'te TTL=3600 saniye ile önbelleklenir. "
         "Bu mekanizma kota aşımını önler ve yanıt süresini iyileştirir."),
        ("KVKK / DR-03, DR-05",
         "Konum verisi yalnızca hizmet sunumu amacıyla, ilgili süre boyunca işlenir. "
         "Kullanıcı veri silme / düzeltme talebinde bulunabilir. Kayıt sırasında "
         "KVKK uyumlu gizlilik bildirimi onaylanır."),
        ("OWASP — SQL Enjeksiyonu",
         "Tüm veritabanı sorguları parametreli (prepared statements) biçimde çalıştırılır."),
        ("OWASP — XSS",
         "React'in varsayılan HTML kaçış mekanizması kullanılır; "
         "dangerouslySetInnerHTML yasaktır."),
        ("OWASP — CSRF",
         "SameSite=Strict cookie politikası uygulanır; "
         "gerekli durumlarda CSRF jeton doğrulama eklenir."),
        ("OWASP — Hız Sınırlama",
         "Giriş uç noktasına dakikada maksimum 10 istek sınırı uygulanır; "
         "aşımda HTTP 429 döndürülür."),
        ("FR-09 — Yetkilendirme (RBAC)",
         "Müşteri yalnızca kendi iş kayıtlarına, Operatör yalnızca atandığı "
         "kayıtlara erişebilir. Yetki dışı erişim HTTP 403 ile reddedilir."),
    ])
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# BÖLÜM 9 — GEREKSİNİM İZLENEBİLİRLİK MATRİSİ
# ──────────────────────────────────────────────────────────────

def section_9(doc):
    h1(doc, "9. Gereksinim İzlenebilirlik Matrisi")
    body(doc,
        "Aşağıdaki matris, SRS.docx belgesindeki 42 gereksinimi (10 UR + 15 FR + "
        "6 DR + 11 NFR) karşılayan tasarım bileşenleri ve önerilen test yöntemleriyle "
        "eşleştirmektedir."
    )

    rows = [
        # ── Kullanıcı Gereksinimleri ──────────────────────────────────────────
        ("UR-01","Y","Müşteri hesabıyla giriş ve talep süreci","AuthModule, AuthService","E2E — Müşteri giriş + talep akışı"),
        ("UR-02","Y","GPS veya harita ile konum belirleme","LocationModule, LocationService","Fonksiyonel — GPS izin + manuel seçim"),
        ("UR-03","Y","Hedef: metin arama veya POI","POIModule, POIService","Entegrasyon — Places API mock testi"),
        ("UR-04","O","Operatörleri mesafe/fiyata göre sıralama","OperatorListModule, OperatorMatchService","Birim — Sıralama algoritma testi"),
        ("UR-05","Y","Operatör seçimi, talep oluşturma, demo ödeme","JobCreateModule, JobService, DemoPaymentService","E2E — Tam iş akışı testi"),
        ("UR-06","Y","Talep durumunu görme","JobStatusModule, JobService","Fonksiyonel — Durum polling / WebSocket"),
        ("UR-07","Y","Operatör profil yönetimi","OperatorDashboardModule, OperatorService","Fonksiyonel — Profil CRUD API testi"),
        ("UR-08","Y","Tarife kurallarını yönetme","OperatorDashboardModule, OperatorService","Birim — Tarife hesaplama doğrulama"),
        ("UR-09","Y","Talep kabul veya red","JobModule, JobService","Entegrasyon — FSM geçiş testi"),
        ("UR-10","Y","İş durumu güncelleme","JobModule, JobService","E2E — Operatör durum güncelleme"),
        # ── Fonksiyonel Gereksinimler ─────────────────────────────────────────
        ("FR-01","Y","Müşteri / operatör kayıt ve giriş akışı","AuthService, AuthModule","Entegrasyon — JWT + rol doğrulama"),
        ("FR-02","Y","İl seçimi ile bağlam başlatma","LocationModule, LocationService","Birim — Şehir dropdown zorunluluğu"),
        ("FR-03","Y","Koordinata göre uygun operatörler","OperatorMatchService, GET /operators/nearby","Entegrasyon — Koordinat filtre API testi"),
        ("FR-04","Y","Her operatör için önizleme tutarı","PricingService, TariffRule","Birim — Tarife formülü doğrulama"),
        ("FR-05","Y","Sunucu taraflı POI sorgusu","POIService, GET /places/search","Entegrasyon — Places API mock + önbellek"),
        ("FR-06","Y","İş talebi oluşturma","JobService, POST /jobs","Entegrasyon — Veritabanı kayıt doğrulama"),
        ("FR-07","Y","Demo ödeme uç noktası","DemoPaymentService, POST /jobs/:id/payment","Birim — Başarı yanıtı doğrulama"),
        ("FR-08","Y","Operatör iş durumu güncelleme","JobService, PATCH /jobs/:id/status","Entegrasyon — FSM geçiş doğrulama"),
        ("FR-09","O","Rol bazlı iş listesi","JobService, GET /jobs","Entegrasyon — Rol erişim kontrol testi"),
        ("FR-10","O","Operatör hizmet alanı","OperatorService, OperatorProfile","Birim — Alan filtresi geometri testi"),
        ("FR-11","O","OpenAPI üzerinden REST endpoint","Express Router, OpenAPI 3.1 spec","Manuel — Swagger UI doğrulama"),
        ("FR-12","O","Oturum zaman aşımında yönlendirme","AuthService, SessionModule","Entegrasyon — Token TTL testi"),
        ("FR-13","O","Geçersiz koordinat için hata kodu","LocationService, zod doğrulama","Birim — Sınır dışı koordinat testi"),
        ("FR-14","O","Eşzamanlı açık talep politikası","JobService (ürün kararı)","Entegrasyon — Eşzamanlı talep senaryosu"),
        ("FR-15","O","Red sonrası alternatif operatör seçimi","JobStatusModule, JobService","E2E — Red + alternatif seçim akışı"),
        # ── Alan Gereksinimleri ───────────────────────────────────────────────
        ("DR-01","—","Tahmini tutar bildirimi","Frontend uyarı bileşeni, PricingService","UI — Tahmini ibare doğrulama"),
        ("DR-02","—","Harita sağlayıcı atıfları","LocationModule (Maps zorunlu attributions)","Manuel — Harita ekranı incelemesi"),
        ("DR-03","—","Konum verisi işlem amacı uyumu","OperatorService, Privacy Notice bileşeni","Manuel — Gizlilik bildirimi kontrolü"),
        ("DR-04","—","Platform hukuki statüsü","Kullanım Şartları (hukuk incelemesi)","Manuel — Hukuk onay süreci"),
        ("DR-05","—","KVKK gizlilik bildirimi","Kayıt akışı — Gizlilik onay adımı","Manuel — KVKK uzman incelemesi"),
        ("DR-06","—","Operatör lisans doğrulaması (kapsam dışı)","Sonraki faza ertelendi","N/A"),
        # ── Fonksiyonel Olmayan Gereksinimler ────────────────────────────────
        ("NFR-01","Y","TLS 1.2+ üretim","Altyapı — HTTPS zorunluluğu, HSTS","Güvenlik — SSL Labs A+ tarama"),
        ("NFR-02","Y","Argon2id parola hash","AuthService — argon2 kütüphanesi","Birim — Hash geri dönüşüm imkânsızlığı"),
        ("NFR-03","Y","API anahtarları sunucu env var'ında","Backend dotenv / AWS Secrets Manager","Güvenlik — Statik kod analizi"),
        ("NFR-04","O","Operatör listesi hızlı yanıt","PricingService, Redis, DB indeksleme","Performans — Yük testi (k6 / Locust)"),
        ("NFR-05","O","Places sorgu önbellekleme","POIService — Redis TTL=3600","Entegrasyon — Önbellek isabet oranı izleme"),
        ("NFR-06","Y","Ana senaryo tek elle tamamlanabilir","Frontend UX, dokunma hedefi ≥44 px","Kullanılabilirlik — Kullanıcı testi oturumu"),
        ("NFR-07","O","ARIA etiketler, renkten bağımsız bilgi","React bileşenleri, aXe denetimi","Erişilebilirlik — Lighthouse / aXe tarama"),
        ("NFR-08","O","OpenAPI sürümlendirme","Express Router /api/v1, OpenAPI 3.1","Manuel — API değişim geri uyumluluğu"),
        ("NFR-09","Y","Hata: genel mesaj + teknik log","Express hata işleyici, loglama ara katmanı","Entegrasyon — HTTP 500 yanıtı kontrolü"),
        ("NFR-10","O","Yatay ölçekleme mimarisi","Durumsuz backend, Redis oturum, ECS Fargate","Performans — Çoklu replika yük testi"),
        ("NFR-11","O","Yedekleme politikası","AWS RDS otomatik yedekleme; RPO/RTO ops tanımı","Manuel — Yedek geri yükleme tatbikatı"),
    ]

    table(doc,
        ["Gereks. ID", "Önc.", "Gereksinim Özeti", "Tasarım Bileşeni", "Test Yöntemi"],
        rows
    )

    p = body(doc,
        f"Toplam: 42 gereksinim  ·  UR: 10  ·  FR: 15  ·  DR: 6  ·  NFR: 11  ·  "
        f"Tarih: {date.today().strftime('%d.%m.%Y')}"
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].italic    = True


# ──────────────────────────────────────────────────────────────
# ANA PROGRAM
# ──────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Sayfa boyutu: Letter — SRS.docx ile aynı (21.59 × 27.94 cm)
    for sec in doc.sections:
        sec.page_width    = Cm(21.59)
        sec.page_height   = Cm(27.94)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # Stilleri SRS.docx ile birebir eşleştir
    configure_styles(doc)

    # Bölümler
    add_cover(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)
    section_6(doc)
    section_7(doc)
    section_8(doc)
    section_9(doc)

    out = r"c:\Users\Lenovo\Desktop\towit\Towit_SDD.docx"
    doc.save(out)
    print(f"OK  Towit_SDD.docx olusturuldu: {out}")


if __name__ == "__main__":
    main()
